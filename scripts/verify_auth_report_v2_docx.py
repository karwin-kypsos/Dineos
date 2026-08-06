"""Verify DineOS_01_Auth_Test_Report_v2.docx contains complete new-field data."""
import json
from pathlib import Path
from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parent.parent
DOCX = PROJECT_ROOT / "docs" / "DineOS_01_Auth_Test_Report_v2.docx"
RESULTS = PROJECT_ROOT / "docs" / "01_auth_v2_results.json"


def full_text(path: Path) -> str:
    doc = Document(str(path))
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def main() -> int:
    if not DOCX.exists():
        print(f"MISSING: {DOCX}")
        return 2
    if not RESULTS.exists():
        print(f"MISSING: {RESULTS}")
        return 2

    text = full_text(DOCX)
    data = json.loads(RESULTS.read_text(encoding="utf-8"))

    print(f"Doc: {DOCX} ({DOCX.stat().st_size:,} bytes)")
    print(f"Total text length: {len(text):,} chars\n")

    checks = []
    # Static markers
    checks += [
        ("Title v2 tag", "Test Run Report (v2)"),
        ("Post-deploy tag", "role_id + role_name"),
        ("Commit hash", "fce96da"),
        ("Server root", "https://dineos-1unt.onrender.com"),
        ("Verification section", "Verification — new fields on live server"),
        ("Changes since v1 heading", "Changes since v1"),
        ("ROLE_METADATA mention", "ROLE_METADATA"),
        ("Pytest total", "62 passed"),
    ]
    # role_id/role_name values from the live responses
    checks += [
        ("Admin role_id in login response",   '"role_id":1,"role_name":"org_admin"'),
        ("Manager role_id in login response", '"role_id":2,"role_name":"manager"'),
        ("Server role_id in login response",  '"role_id":3,"role_name":"server"'),
        ("Cashier role_id in login response", '"role_id":4,"role_name":"cashier"'),
        ("Me role_id in Me response",         '"role_id":1'),
        ("Me role_name in Me response",       '"role_name":"org_admin"'),
    ]
    # Check that every request's response body (from the JSON results) is present in the docx
    for r in data["results"]:
        body = r["responseBody"]
        # 204 responses are represented by the "(empty — 204 No Content)" placeholder
        if "204 No Content" in body:
            snippet = "204 No Content"
        else:
            # Use a distinctive substring from the middle of each body
            snippet = body[:100].replace("\n", " ")
        checks.append((f"Response body #{r['step']} ({r['name']}) present", snippet))

    missing = []
    for label, needle in checks:
        ok = needle in text
        print(f"[{'OK  ' if ok else 'MISS'}] {label}")
        if not ok:
            missing.append((label, needle))

    print()
    if missing:
        print(f"FAILED: {len(missing)} of {len(checks)} checks missing.")
        for label, needle in missing[:6]:
            print(f"  - {label}: {needle[:120]}")
        return 1
    print(f"SUCCESS: all {len(checks)} checks present.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
