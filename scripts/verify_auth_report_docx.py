"""Verify the generated docx contains every request/response in full."""
from pathlib import Path
from docx import Document

DOCX = Path(__file__).resolve().parent.parent / "docs" / "DineOS_01_Auth_Test_Report.docx"

# Signature substrings that MUST appear somewhere in the doc.
# JWT signatures are unique per token (last dot-segment), so their presence
# proves the whole token is intact.
CHECKS = [
    # --- Step 1: Login — Admin ---
    ("1. Login — Admin (heading)", "1. Login — Admin"),
    ("Admin login request email", '"email": "admin@demo-bistro.demo"'),
    ("Admin login refresh JWT sig", "384n9AQhwTVBQoYCm7l3K8d4aiqDjzVU-AXHzcqEKcM"),
    ("Admin login access JWT sig",  "AUxmoelkd8Sp34zi4QelMWPuymf-mZmoZIL5BPyXeCY"),

    # --- Step 2: Login — Manager ---
    ("2. Login — Manager (heading)", "2. Login — Manager"),
    ("Manager login request email", '"email": "manager@demo-bistro.demo"'),
    ("Manager refresh JWT sig", "yOFvA6chP0oBOyaBgD8VYtyWeBO0LubQ5ahkU73pNjc"),
    ("Manager access JWT sig",  "xGJSf_f07e662TJ1zO2B56K6UtnEPlVAipSbj5oOD4w"),

    # --- Step 3: Login — Server ---
    ("3. Login — Server (heading)", "3. Login — Server"),
    ("Server login request email", '"email": "server@demo-bistro.demo"'),
    ("Server refresh JWT sig", "7n_jKDUIcHaeTLFKjYgX6KVKXFrxex0epsCIMQDoS5I"),
    ("Server access JWT sig",  "LM0DXTY4TxXKMpdkLj0RkzqVn4khICsDeQMaG5T3SmE"),

    # --- Step 4: Login — Cashier ---
    ("4. Login — Cashier (heading)", "4. Login — Cashier"),
    ("Cashier login request email", '"email": "cashier@demo-bistro.demo"'),
    ("Cashier refresh JWT sig", "LFZGWLqsK1p12eik3Legg9x3bj8jMYTkzhfDrk4IEcw"),
    ("Cashier access JWT sig",  "7-O0CPvs2FWtdOrOgAViWkmbRG2UhVyQTiYk1W3FDq8"),

    # --- Step 5: Me ---
    ("5. Me (heading)", "5. Me"),
    ("Me user_id", "acc5014f-08e7-4c3a-b2fc-0ca2298cc677"),
    ("Me restaurant slug", '"slug":"demo-bistro"'),
    ("Me feature flag: kitchen", '"kitchen_enabled":true'),
    ("Me feature flag: billing", '"billing_enabled":true'),
    ("Me feature flag: notifications", '"notifications_enabled":true'),
    ("Me feature flag: realtime", '"realtime_enabled":true'),
    ("Me created_at", "2026-08-03T12:04:47.278027Z"),

    # --- Step 6: Refresh Token ---
    ("6. Refresh Token (heading)", "6. Refresh Token"),
    ("Refresh request uses original admin refresh",
     "384n9AQhwTVBQoYCm7l3K8d4aiqDjzVU-AXHzcqEKcM"),
    ("Refresh response new access sig",
     "6mfIqoHVY80ibtzVt89YrKilLd94NtT-bNRomhM5Wck"),
    ("Refresh response new refresh sig",
     "pSjaK_zS-FdvTB6s8uWnygepcV3vNJdwQ0Q4FL5lNR4"),

    # --- Step 7: Change Password ---
    ("7. Change Password (heading)", "7. Change Password"),
    ("Change pwd request current",  '"current_password": "Demo@1234"'),
    ("Change pwd request new",      '"new_password": "Demo@1234"'),
    ("Change pwd response",         '"detail":"Password changed."'),

    # --- Step 8: Forgot Password ---
    ("8. Forgot Password (heading)", "8. Forgot Password"),
    ("Forgot pwd request email",    '"email": "admin@demo-bistro.demo"'),
    ("Forgot pwd generic ack",      "If that email is registered"),

    # --- Step 9: Logout ---
    ("9. Logout (heading)", "9. Logout"),
    ("Logout request uses rotated refresh",
     "pSjaK_zS-FdvTB6s8uWnygepcV3vNJdwQ0Q4FL5lNR4"),
    ("Logout 204 marker", "204 No Content"),

    # --- Meta ---
    ("Server root updated to Render", "https://dineos-1unt.onrender.com"),
    ("Tenant slug present",  "demo-bistro"),
    ("Collection name",      "DineOS v2 Backend API (Phase 1)"),
    ("Folder name",          "01 Auth"),
]


def full_text(path: Path) -> str:
    """Concatenate every paragraph AND every table cell in the doc."""
    doc = Document(str(path))
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def main() -> int:
    if not DOCX.exists():
        print(f"MISSING: {DOCX}")
        return 2
    text = full_text(DOCX)
    print(f"Doc: {DOCX} ({DOCX.stat().st_size:,} bytes)")
    print(f"Total text length: {len(text):,} chars\n")

    missing = []
    for label, needle in CHECKS:
        ok = needle in text
        status = "OK  " if ok else "MISS"
        print(f"[{status}] {label}")
        if not ok:
            missing.append((label, needle))

    print()
    if missing:
        print(f"FAILED: {len(missing)} of {len(CHECKS)} checks missing.")
        for label, needle in missing:
            print(f"  - {label}: expected substring {needle!r}")
        return 1
    print(f"SUCCESS: all {len(CHECKS)} checks present.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
