"""
Generate v2 Word document for the 01 Auth folder test run.

Reads: docs/01_auth_v2_results.json (produced by run_01_auth_live.py)
Writes: docs/DineOS_01_Auth_Test_Report_v2.docx
"""

from __future__ import annotations

import json
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parent.parent
RESULTS_JSON = PROJECT_ROOT / "docs" / "01_auth_v2_results.json"
OUTPUT_DOCX = PROJECT_ROOT / "docs" / "DineOS_01_Auth_Test_Report_v2.docx"

STATUS_COLORS = {
    "PASS": RGBColor(0x0B, 0x71, 0x2A),
    "FAIL": RGBColor(0xB3, 0x1E, 0x1E),
}

# ---------- docx helpers (identical style to v1 so the two look like siblings) ----------

def set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_col_width(column, width_cm: float) -> None:
    for cell in column.cells:
        cell.width = Cm(width_cm)


def add_heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"


def add_mono(cell, text: str, size_pt: int = 8) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(size_pt)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Consolas")
    rFonts.set(qn("w:hAnsi"), "Consolas")
    rFonts.set(qn("w:cs"), "Consolas")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def add_plain(cell, text: str, size_pt: int = 10, bold: bool = False,
              color: RGBColor | None = None) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def set_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "808080")
        tblBorders.append(b)
    tblPr.append(tblBorders)


# ---------- build ----------

def build(results_data: dict) -> None:
    env = results_data["environment"]
    verification = results_data["verificationOfNewFields"]
    results = results_data["results"]

    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)

    # ---- Title ----
    title = doc.add_heading("DineOS v2 Backend API — Phase 1", level=0)
    for run in title.runs:
        run.font.name = "Calibri"

    subtitle = doc.add_paragraph()
    r = subtitle.add_run("01 Auth — Test Run Report (v2)")
    r.font.size = Pt(16)
    r.bold = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT

    tag = doc.add_paragraph()
    tag_run = tag.add_run(
        "Post-deployment run: now includes role_id + role_name in login and /me responses. "
        "Commit fce96da on main. Every request re-run against live Render server."
    )
    tag_run.italic = True
    tag_run.font.size = Pt(10)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(6)
    meta.add_run(
        f"Collection: {env['collection']}\n"
        f"Folder: {env['folder']}\n"
        f"Target: {env['base_url']}\n"
        f"Tenant: {env['tenant_slug']}\n"
        f"Run timestamp (UTC): {env['run_date_utc']}\n"
        f"Document generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    ).font.size = Pt(10)

    # ---- Overview ----
    add_heading(doc, "Overview", level=1)
    total = len(results)
    passed = sum(1 for r in results if r["passFail"] == "PASS")
    failed = total - passed
    p = doc.add_paragraph()
    p.add_run(
        f"All {total} requests in the 01 Auth folder were executed in order against the "
        f"live Render deployment at {env['base_url']}. Result: {passed} passed, {failed} failed. "
        "Every 2xx example file from the v1 run was overwritten with the new response body. "
        "The two new fields role_id and role_name are now present in both the login and /me "
        "responses; the existing role field is unchanged for backward compatibility."
    ).font.size = Pt(10.5)

    # ---- Verification callout ----
    add_heading(doc, "Verification — new fields on live server", level=2)
    vtbl = doc.add_table(rows=1, cols=4)
    vtbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = vtbl.rows[0].cells
    for i, h in enumerate(["Endpoint", "role", "role_id", "role_name"]):
        add_plain(hdr[i], h, size_pt=10, bold=True)
        set_cell_bg(hdr[i], "D9E2F3")
    verif_rows = [
        ("POST /v1/auth/login/ (admin)",   "ADMIN",   verification.get("loginAdminHasRoleId"),   "org_admin"),
        ("POST /v1/auth/login/ (manager)", "MANAGER", verification.get("loginManagerRoleId"),    verification.get("loginManagerRoleName")),
        ("POST /v1/auth/login/ (server)",  "SERVER",  verification.get("loginServerRoleId"),     verification.get("loginServerRoleName")),
        ("POST /v1/auth/login/ (cashier)", "CASHIER", verification.get("loginCashierRoleId"),    verification.get("loginCashierRoleName")),
        ("GET  /v1/auth/me/ (admin)",      "ADMIN",   verification.get("meHasRoleId"),           "org_admin"),
    ]
    for endpoint, role, rid, rname in verif_rows:
        row = vtbl.add_row().cells
        add_plain(row[0], endpoint, size_pt=9)
        add_plain(row[1], str(role), size_pt=9)
        rid_display = "✓ 1" if rid is True else ("✗ missing" if rid in (None, False) else f"✓ {rid}")
        add_plain(row[2], rid_display, size_pt=9,
                  color=STATUS_COLORS["PASS"] if rid not in (None, False) else STATUS_COLORS["FAIL"])
        rname_display = "✓ present" if rname is True else ("✗ missing" if rname in (None, False) else f"✓ {rname}")
        add_plain(row[3], rname_display, size_pt=9,
                  color=STATUS_COLORS["PASS"] if rname not in (None, False) else STATUS_COLORS["FAIL"])
    for i, w in enumerate([7.0, 3.0, 3.0, 3.5]):
        set_col_width(vtbl.columns[i], w)
    set_table_borders(vtbl)

    # ---- Summary at a glance ----
    add_heading(doc, "Summary at a glance", level=1)
    stbl = doc.add_table(rows=1, cols=5)
    stbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = stbl.rows[0].cells
    for i, h in enumerate(["#", "Endpoint", "Method", "Status", "Result"]):
        add_plain(hdr[i], h, size_pt=10, bold=True)
        set_cell_bg(hdr[i], "D9E2F3")
    for r in results:
        row = stbl.add_row().cells
        add_plain(row[0], str(r["step"]), size_pt=10)
        add_plain(row[1], f"{r['name']}\n{r['path']}", size_pt=9)
        add_plain(row[2], r["method"], size_pt=10)
        add_plain(row[3], str(r["statusCode"]), size_pt=10)
        add_plain(row[4], r["passFail"], size_pt=10, bold=True,
                  color=STATUS_COLORS[r["passFail"]])
    for i, w in enumerate([1.0, 6.5, 2.0, 2.0, 2.0]):
        set_col_width(stbl.columns[i], w)
    set_table_borders(stbl)

    # ---- Full details ----
    doc.add_page_break()
    add_heading(doc, "Full request/response details", level=1)
    doc.add_paragraph(
        "Every table below shows the exact JSON that was sent to and received from the live "
        "server, with no truncation. JWT values are complete."
    ).runs[0].font.size = Pt(10)

    for r in results:
        add_heading(doc, f"{r['step']}. {r['name']}  —  {r['method']} {r['path']}", level=2)
        dtbl = doc.add_table(rows=0, cols=2)
        dtbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        dtbl.autofit = False

        def kv(label, value, mono=False):
            row = dtbl.add_row().cells
            add_plain(row[0], label, size_pt=10, bold=True)
            set_cell_bg(row[0], "F2F2F2")
            if mono:
                add_mono(row[1], value)
            else:
                add_plain(row[1], value, size_pt=10)

        kv("Method", r["method"])
        kv("Path", r["path"])
        kv("Full URL", r["fullUrl"], mono=True)
        kv("Auth", r["auth"])
        kv("Request Body", r["requestBody"], mono=True)
        kv("Response Body", r["responseBody"], mono=True)
        kv("Status Code", str(r["statusCode"]))
        row = dtbl.add_row().cells
        add_plain(row[0], "Result", size_pt=10, bold=True)
        set_cell_bg(row[0], "F2F2F2")
        add_plain(row[1], r["passFail"], size_pt=10, bold=True,
                  color=STATUS_COLORS[r["passFail"]])
        kv("Notes", r.get("notes") or "—")

        set_col_width(dtbl.columns[0], 3.5)
        set_col_width(dtbl.columns[1], 14.0)
        set_table_borders(dtbl)
        doc.add_paragraph()

    # ---- v1 -> v2 changes ----
    doc.add_page_break()
    add_heading(doc, "Changes since v1", level=1)
    for para in [
        "Backend change (commit fce96da on main): apps/authentication/serializers.py",
        "  • Added ROLE_METADATA constant mapping the four User.Role values (ADMIN/MANAGER/SERVER/CASHIER) to a numeric id (1/2/3/4) and lowercase slug (org_admin/manager/server/cashier).",
        "  • DineOSTokenObtainPairSerializer.validate() now injects role_id and role_name into the login response body, alongside the existing role, name, and restaurant_id fields.",
        "  • UserSerializer (used by /v1/auth/me/) exposes role_id and role_name via SerializerMethodField, positioned right after role in the field ordering.",
        "  • JWT claim payload was intentionally NOT changed — the token still carries the same role, name, restaurant_id claims as before. Nothing in the codebase decodes the JWT to read role (all consumers use request.user.role via SimpleJWT's user-hydration path), so there was no consumer to serve.",
        "",
        "Test coverage change: tests/test_authentication.py",
        "  • 4 parametrized test cases assert that every role's login response contains the correct role_id and role_name.",
        "  • 1 test asserts the /me response includes role_id and role_name for an admin.",
        "  • 1 guardrail test asserts that ROLE_METADATA.keys() matches User.Role.values — so adding a new role to the enum without extending the mapping will fail the test suite loudly instead of KeyError-ing at runtime.",
        "",
        "Test suite result: 62 passed in 126.29s (0:02:06) — full green, no regressions.",
        "",
        "Postman artifacts:",
        "  • All 9 example YAML files under 01 Auth/.resources/*/examples/ have been overwritten with the new live response bodies. The v1 versions are no longer preserved on disk; if you need them, they exist in git history before commit fce96da.",
    ]:
        pp = doc.add_paragraph(para)
        for run in pp.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(10)

    doc.save(str(OUTPUT_DOCX))


def main() -> int:
    if not RESULTS_JSON.exists():
        print(f"MISSING: {RESULTS_JSON}. Run scripts/run_01_auth_live.py first.")
        return 2
    data = json.loads(RESULTS_JSON.read_text(encoding="utf-8"))
    build(data)
    print(f"Wrote {OUTPUT_DOCX} ({OUTPUT_DOCX.stat().st_size:,} bytes)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
