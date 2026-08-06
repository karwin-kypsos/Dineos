"""
Append a folder's live-test results as a new section onto the existing
master report docx (docs/DineOS_01_Auth_Test_Report_v2.docx), which grows
into the running multi-folder test report as each folder is completed.

Usage:
    .venv\\Scripts\\python.exe scripts\\append_folder_report.py 02_staff
"""

from __future__ import annotations

import argparse
import json
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parent.parent
MASTER_DOCX = PROJECT_ROOT / "docs" / "DineOS_01_Auth_Test_Report_v2.docx"

STATUS_COLORS = {
    "PASS": RGBColor(0x0B, 0x71, 0x2A),
    "FAIL": RGBColor(0xB3, 0x1E, 0x1E),
}


def set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_col_width(column, width_cm: float) -> None:
    for cell in column.cells:
        cell.width = Cm(width_cm)


def add_heading(doc: Document, text: str, level: int):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
    return h


def add_mono(cell, text: str, size_pt: int = 8) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(size_pt)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Consolas")
    rFonts.set(qn("w:hAnsi"), "Consolas")
    rFonts.set(qn("w:cs"), "Consolas")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def add_plain(cell, text: str, size_pt: int = 10, bold: bool = False,
              color: RGBColor | None = None) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def set_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "808080")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def append_section(doc: Document, results_data: dict) -> None:
    env = results_data["environment"]
    results = results_data["results"]
    findings = results_data.get("findings", [])

    doc.add_page_break()
    add_heading(doc, f"{env['folder']} — Live Test Run Report (v2)", level=1)

    tag = doc.add_paragraph()
    tag_run = tag.add_run(
        f"Every request executed live against the deployed Render server via "
        f"{env.get('runner', 'Newman (Postman CLI runner)')}. No mocked or fabricated data."
    )
    tag_run.italic = True
    tag_run.font.size = Pt(10)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(6)
    meta.add_run(
        f"Collection: {env['collection']}\n"
        f"Folder: {env['folder']}\n"
        f"Target: {env['base_url']}\n"
        f"Tenant: {env['tenant_slug']}\n"
        f"Run timestamp (UTC): {env['run_date_utc']}\n"
        f"Document generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    ).font.size = Pt(10)

    add_heading(doc, "Overview", level=2)
    total = len(results)
    passed = sum(1 for r in results if r["passFail"] == "PASS")
    failed = total - passed
    p = doc.add_paragraph()
    p.add_run(
        f"All {total} requests in the {env['folder']} folder were executed in order against the "
        f"live Render deployment at {env['base_url']}. Result: {passed} passed, {failed} failed."
    ).font.size = Pt(10.5)
    if findings:
        p2 = doc.add_paragraph()
        p2.add_run(
            f"{len(findings)} finding(s) were surfaced during the run — see 'Findings' subsection "
            "below for details."
        ).font.size = Pt(10.5)

    add_heading(doc, "Summary at a glance", level=2)
    stbl = doc.add_table(rows=1, cols=5)
    stbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = stbl.rows[0].cells
    for i, h in enumerate(["#", "Endpoint", "Method", "Status", "Result"]):
        add_plain(hdr[i], h, size_pt=10, bold=True)
        set_cell_bg(hdr[i], "D9E2F3")
    for r in results:
        row = stbl.add_row().cells
        add_plain(row[0], str(r["step"]), size_pt=10)
        add_plain(row[1], f"{r['name']}\n{r['path']}", size_pt=9)
        add_plain(row[2], r["method"], size_pt=10)
        add_plain(row[3], str(r["statusCode"]), size_pt=10)
        add_plain(row[4], r["passFail"], size_pt=10, bold=True,
                  color=STATUS_COLORS[r["passFail"]])
    for i, w in enumerate([1.0, 6.5, 2.0, 2.0, 2.0]):
        set_col_width(stbl.columns[i], w)
    set_table_borders(stbl)

    add_heading(doc, "Full request/response details", level=2)
    doc.add_paragraph(
        "Every table below shows the exact JSON that was sent to and received from the live "
        "server, with no truncation. JWT values are complete."
    ).runs[0].font.size = Pt(10)

    for r in results:
        add_heading(doc, f"{r['step']}. {r['name']}  —  {r['method']} {r['path']}", level=3)
        dtbl = doc.add_table(rows=0, cols=2)
        dtbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        dtbl.autofit = False

        def kv(label, value, mono=False):
            row = dtbl.add_row().cells
            add_plain(row[0], label, size_pt=10, bold=True)
            set_cell_bg(row[0], "F2F2F2")
            if mono:
                add_mono(row[1], value)
            else:
                add_plain(row[1], value, size_pt=10)

        kv("Method", r["method"])
        kv("Path", r["path"])
        kv("Full URL", r["fullUrl"], mono=True)
        kv("Auth", r["auth"])
        kv("Request Body", r["requestBody"], mono=True)
        kv("Response Body", r["responseBody"], mono=True)
        kv("Status Code", str(r["statusCode"]))
        row = dtbl.add_row().cells
        add_plain(row[0], "Result", size_pt=10, bold=True)
        set_cell_bg(row[0], "F2F2F2")
        add_plain(row[1], r["passFail"], size_pt=10, bold=True,
                  color=STATUS_COLORS[r["passFail"]])
        kv("Notes", r.get("notes") or "—")

        set_col_width(dtbl.columns[0], 3.5)
        set_col_width(dtbl.columns[1], 14.0)
        set_table_borders(dtbl)
        doc.add_paragraph()

    if findings:
        add_heading(doc, "Findings", level=2)
        for f in findings:
            add_heading(doc, f["title"], level=3)
            sev = doc.add_paragraph()
            sev.add_run(f"Severity: {f['severity']}").bold = True
            sev.runs[0].font.size = Pt(10)
            desc = doc.add_paragraph(f["description"])
            desc.runs[0].font.size = Pt(10)
            ev = f.get("evidence")
            if ev:
                etbl = doc.add_table(rows=0, cols=2)
                etbl.alignment = WD_TABLE_ALIGNMENT.LEFT

                def kv2(label, value):
                    row = etbl.add_row().cells
                    add_plain(row[0], label, size_pt=10, bold=True)
                    set_cell_bg(row[0], "F2F2F2")
                    add_mono(row[1], str(value))

                for k, v in ev.items():
                    kv2(k, v)
                set_col_width(etbl.columns[0], 3.5)
                set_col_width(etbl.columns[1], 14.0)
                set_table_borders(etbl)
            doc.add_paragraph()


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("slug", help="Results file slug, e.g. 02_staff (reads docs/02_staff_v2_results.json)")
    args = ap.parse_args()

    results_json = PROJECT_ROOT / "docs" / f"{args.slug}_v2_results.json"
    if not results_json.exists():
        print(f"MISSING: {results_json}")
        return 2
    data = json.loads(results_json.read_text(encoding="utf-8"))

    doc = Document(str(MASTER_DOCX))
    append_section(doc, data)
    doc.save(str(MASTER_DOCX))
    print(f"Appended {data['environment']['folder']} section -> {MASTER_DOCX} ({MASTER_DOCX.stat().st_size:,} bytes)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
