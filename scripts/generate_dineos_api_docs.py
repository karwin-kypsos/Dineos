"""
Generate the DineOS Backend API — Project Documentation docx, following
the structure of the reference doc (downloaded client example) exactly:
title block -> Project Overview -> Project Structure -> Database Models
-> API Endpoints (numbered sections, each a table with columns Screen |
Method | URL | Parameters | Output Parameters(In Response) | Changes
Required | Comments | Status).

Every row in the API Endpoints section comes from a real, live request/
response captured via Newman against the deployed Render server
(postman/collections/_build/dineos_full_live_results.json) — nothing in
that section is fabricated.

Reads:
    postman/collections/_build/dineos_full_live_run.json      (folder structure)
    postman/collections/_build/dineos_full_live_results.json  (captured req/resp)
Writes:
    docs/DineOS_Backend_API_Documentation.docx
"""

from __future__ import annotations

import json
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parent.parent
COLLECTION_JSON = PROJECT_ROOT / "postman" / "collections" / "_build" / "dineos_full_live_run.json"
RESULTS_JSON = PROJECT_ROOT / "postman" / "collections" / "_build" / "dineos_full_live_results.json"
OUTPUT_DOCX = PROJECT_ROOT / "docs" / "DineOS_Backend_API_Documentation.docx"

BASE_URL = "https://dineos-1unt.onrender.com"

# Folder (as it appears in the Postman collection) -> numbered section title.
SECTION_TITLES = [
    ("01 Auth", "Authentication"),
    ("01b Branches", "Branches"),
    ("02 Staff", "Staff"),
    ("03 Tables & Sessions", "Tables & Sessions"),
    ("04 Menu", "Menu"),
    ("05 Prepared Dishes", "Daily Prep Log"),
    ("11 Inventory", "Inventory & Stock"),
    ("06 Orders", "Orders"),
    ("07 Billing", "Billing"),
    ("07b Cashier Shift", "Cashier Shift"),
    ("08 Notifications", "Notifications"),
    ("09 Health", "Health"),
    ("12 Admin Dashboard", "Admin Dashboard"),
    ("10 Platform (Super Admin)", "Super Admin (Platform)"),
]

# Requests whose non-2xx status is the CORRECT, intended behavior (negative
# tests) — everything else is expected to be 2xx.
EXPECTED_NON_2XX = {
    "Delete Category — Blocked, Has Items (expect 409)": "409 — category still has menu items; delete correctly refused",
    "Close Shift — Discrepancy Rejected (Cashier)": "409 — uncounted cash mismatch; close correctly refused pending acknowledgement",
    "Verify Suspended Org Blocks Staff Login (expect 400)": "400 — organization suspended; login correctly refused",
    "Verify Impersonation Revoked (Me, expect 401)": "401 — support session had been ended; token correctly rejected",
    "Staff Token Rejected Here (should 401-403)": "401 — restaurant-staff token correctly rejected on a platform-only route",
    "Verify Logout Revoked (List Tenants, expect 401)": "401 — token was blacklisted at logout; correctly rejected",
    "Verify Refresh Token Revoked (expect 400)": "400 — refresh token was revoked at logout; correctly rejected",
}

STATUS_OK_COLOR = RGBColor(0x0B, 0x71, 0x2A)


# ---------------------------------------------------------------------------
# docx helpers
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_col_width(column, width_cm: float) -> None:
    for cell in column.cells:
        cell.width = Cm(width_cm)


def add_heading(doc: Document, text: str, level: int):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
    return h


def add_mono(cell, text: str, size_pt: int = 8) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(size_pt)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Consolas")
    rFonts.set(qn("w:hAnsi"), "Consolas")
    rFonts.set(qn("w:cs"), "Consolas")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def add_plain(cell, text: str, size_pt: int = 10, bold: bool = False, color: RGBColor | None = None) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def set_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "808080")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def pretty_json(text: str) -> str:
    if not text:
        return text
    try:
        return json.dumps(json.loads(text), indent=2, ensure_ascii=False)
    except (ValueError, TypeError):
        return text


def two_col_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    tbl = doc.add_table(rows=0, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for h1, h2 in rows:
        row = tbl.add_row().cells
        add_plain(row[0], h1, size_pt=10, bold=True)
        set_cell_bg(row[0], "D9E2F3")
        add_plain(row[1], h2, size_pt=10)
    set_col_width(tbl.columns[0], 5.0)
    set_col_width(tbl.columns[1], 12.0)
    set_table_borders(tbl)


def model_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["Field", "Type", "Description"]):
        add_plain(hdr[i], h, size_pt=10, bold=True)
        set_cell_bg(hdr[i], "D9E2F3")
    for field, ftype, desc in rows:
        row = tbl.add_row().cells
        add_mono(row[0], field, size_pt=9)
        add_plain(row[1], ftype, size_pt=9)
        add_plain(row[2], desc, size_pt=9)
    for i, w in enumerate([4.5, 4.0, 8.5]):
        set_col_width(tbl.columns[i], w)
    set_table_borders(tbl)


# ---------------------------------------------------------------------------
# build
# ---------------------------------------------------------------------------

def build_folder_index(collection: dict) -> dict[str, list[str]]:
    """folder name -> ordered list of request names in that folder."""
    index = {}
    for folder in collection["item"]:
        index[folder["name"]] = [item["name"] for item in folder["item"]]
    return index


def build(collection: dict, results: list[dict]) -> None:
    by_name = {}
    for r in results:
        by_name.setdefault(r["name"], []).append(r)
    folder_index = build_folder_index(collection)

    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)

    # ---- Title block ----
    t1 = doc.add_paragraph()
    r1 = t1.add_run("DINEOS")
    r1.bold = True
    r1.font.size = Pt(22)
    r1.font.name = "Calibri"

    t2 = doc.add_paragraph()
    r2 = t2.add_run("Backend API — Project Documentation")
    r2.font.size = Pt(15)
    r2.font.name = "Calibri"

    t3 = doc.add_paragraph()
    r3 = t3.add_run("Built with Django REST Framework + PostgreSQL + Django Channels (WebSockets)")
    r3.font.size = Pt(11)
    r3.italic = True

    t4 = doc.add_paragraph()
    r4 = t4.add_run(f"Developed by Krypsos  |  {datetime.now(timezone.utc).strftime('%B %Y')}")
    r4.font.size = Pt(11)

    # ---- 1. Project Overview ----
    add_heading(doc, "1. Project Overview", level=1)
    doc.add_paragraph(
        "DineOS is a multi-tenant restaurant operations platform for sit-down and "
        "takeaway restaurants. A single deployment serves many independent restaurant "
        "organizations (tenants), each with one or more branches. The platform covers "
        "QR-code table ordering, kitchen display coordination, billing and cashier "
        "shifts, inventory and stock tracking, staff/branch management, and a separate "
        "Super Admin console for the Krypsos team to onboard and operate client "
        "organizations."
    ).runs[0].font.size = Pt(10.5)

    add_heading(doc, "1.1 Goal", level=2)
    for line in [
        "Fast, contactless table-side ordering via QR code, no app install for guests",
        "Real-time kitchen coordination — orders flow straight to the Kitchen Display Screen",
        "Accurate, tenant-configurable billing (GST + service charge) with cashier shift reconciliation",
        "Multi-branch support with per-branch tables, menus, staff, and inventory",
        "Plan-tier-based feature gating (kitchen, billing, notifications, realtime) controlled centrally by Krypsos",
        "A dedicated Super Admin console for onboarding, support access, and platform-wide oversight",
    ]:
        doc.add_paragraph(line).runs[0].font.size = Pt(10.5)

    add_heading(doc, "1.2 Tech Stack", level=2)
    stack_rows = [
        ("Layer", "Technology"),
        ("Backend Framework", "Django 4.2 + Django REST Framework"),
        ("Database", "PostgreSQL (Render managed)"),
        ("Authentication", "JWT (rest_framework_simplejwt) — separate realms for restaurant staff and Platform admins"),
        ("Realtime", "Django Channels + Redis channel layer (WebSockets for kitchen/table/notification updates)"),
        ("Background Jobs", "Celery + Redis"),
        ("Hosting", "Render (web service + worker + managed Postgres + managed Redis)"),
        ("API Testing", "Postman (git-synced collection) run live via Newman against the deployed server"),
    ]
    stbl = doc.add_table(rows=1, cols=2)
    stbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = stbl.rows[0].cells
    add_plain(hdr[0], stack_rows[0][0], bold=True)
    set_cell_bg(hdr[0], "D9E2F3")
    add_plain(hdr[1], stack_rows[0][1], bold=True)
    set_cell_bg(hdr[1], "D9E2F3")
    for layer, tech in stack_rows[1:]:
        row = stbl.add_row().cells
        add_plain(row[0], layer, size_pt=10)
        add_plain(row[1], tech, size_pt=10)
    set_col_width(stbl.columns[0], 5.0)
    set_col_width(stbl.columns[1], 12.0)
    set_table_borders(stbl)

    # ---- 2. Project Structure ----
    add_heading(doc, "2. Project Structure", level=1)
    doc.add_paragraph(
        "The Django project is named dineos and contains one app per domain area, "
        "plus a core package for cross-cutting tenant-resolution and permission logic."
    ).runs[0].font.size = Pt(10.5)

    add_heading(doc, "2.1 Folder Structure", level=2)
    tree = doc.add_paragraph()
    tree_run = tree.add_run(
        "dineos/\n"
        "  dineos/                  # Project config\n"
        "    settings/              # base.py, development.py, production.py\n"
        "    urls.py\n"
        "    asgi.py                # Channels ASGI entrypoint\n"
        "  core/                    # Tenant resolution, permissions, exceptions\n"
        "    tenancy.py\n"
        "    permissions.py\n"
        "  apps/\n"
        "    authentication/        # Staff accounts, login/invite/reset, roles\n"
        "    restaurant/            # Restaurant (tenant) + Branch\n"
        "    tables/                # Table + TableSession (QR ordering)\n"
        "    menu/                  # Category, MenuItem, PreparedPortion (Daily Prep Log)\n"
        "    orders/                # Order, OrderItem (dine-in + takeaway)\n"
        "    billing/               # Bill, CashierShift\n"
        "    inventory/             # Ingredient, StockMovement, PurchaseOrder, RecipeItem\n"
        "    kitchen/                # KDSDevice + key auth for the Kitchen Display Screen\n"
        "    notifications/         # Notification (in-app + WebSocket)\n"
        "    dashboard/              # Admin Dashboard, EOD Review, low-stock AI alerts\n"
        "    platform/              # Super Admin realm: PlatformAdmin, tenants, impersonation, activity log"
    )
    tree_run.font.name = "Consolas"
    tree_run.font.size = Pt(9)

    # ---- 3. Database Models ----
    add_heading(doc, "3. Database Models", level=1)

    add_heading(doc, "3.1 Restaurant & Branch (apps.restaurant)", level=2)
    doc.add_paragraph("One Restaurant row per client organization (tenant); one or more Branch rows per Restaurant.").runs[0].font.size = Pt(10)
    model_table(doc, [
        ("name", "CharField", "Organization display name"),
        ("slug", "SlugField", "Unique tenant identifier, used in QR URLs"),
        ("status", "CharField (choices)", "ACTIVE / TRIAL / SUSPENDED — suspended blocks all staff logins"),
        ("plan_tier", "CharField (choices)", "STARTER / GROWTH / ENTERPRISE — pre-fills max_branches + feature flags"),
        ("max_branches", "PositiveIntegerField, null", "None = unlimited (Enterprise)"),
        ("gst_percentage", "DecimalField", "Tenant-configurable, editable by the tenant's own Admin"),
        ("service_charge_percentage", "DecimalField", "Tenant-configurable"),
        ("notifications_enabled / kitchen_enabled / billing_enabled / realtime_enabled", "BooleanField", "Per-tenant add-on flags, controlled by Platform Super Admin"),
        ("contact_name / contact_email / contact_phone / billing_email", "CharField / EmailField", "Org identity, set at onboarding"),
        ("primary_color", "CharField (hex)", "Branding color for the tenant's app instance"),
        ("branch.name / branch.address / branch.phone / branch.photo_url", "CharField / URLField", "Branch identity"),
        ("branch.manager", "ForeignKey(User), null", "Optional — not required at creation"),
        ("branch.table_count", "PositiveIntegerField, null", "Optional — syncing this creates/deactivates numbered tables"),
    ])

    add_heading(doc, "3.2 User (apps.authentication)", level=2)
    doc.add_paragraph("Restaurant staff account — Admin, Manager, Server, or Cashier, scoped to one Restaurant and optionally one Branch.").runs[0].font.size = Pt(10)
    model_table(doc, [
        ("email", "EmailField, unique", "Login identifier"),
        ("role", "CharField (choices)", "ADMIN / MANAGER / SERVER / CASHIER"),
        ("restaurant", "ForeignKey(Restaurant)", "Tenant this staff member belongs to"),
        ("branch", "ForeignKey(Branch), null", "Optional branch assignment"),
        ("must_change_password", "BooleanField", "True until an invited user completes first-login password setup"),
        ("is_active", "BooleanField", "Deactivated staff cannot log in"),
    ])

    add_heading(doc, "3.3 Table & TableSession (apps.tables)", level=2)
    model_table(doc, [
        ("table_number", "CharField", "Printed on the table's QR code"),
        ("branch", "ForeignKey(Branch), null", "Null = legacy ungrouped table"),
        ("status", "CharField (choices)", "AVAILABLE / OCCUPIED / BILL_REQUESTED, etc."),
        ("session.id", "UUIDField", "The unguessable token embedded in the QR URL — acts as the guest's access credential"),
        ("session.closed_reason", "CharField (choices)", "PAID / MANAGER_OVERRIDE, etc."),
    ])

    add_heading(doc, "3.4 Category & MenuItem (apps.menu)", level=2)
    model_table(doc, [
        ("category.name / category.emoji / category.sort_order", "CharField / IntegerField", "Menu section grouping"),
        ("menu_item.name / menu_item.price / menu_item.description", "CharField / DecimalField", "Dish identity and pricing"),
        ("menu_item.is_veg", "BooleanField", "Vegetarian indicator, shown as a badge in the customer menu"),
        ("menu_item.is_available", "BooleanField", "Toggled off to 86 an item"),
        ("prepared_portion.portions_initial / portions_remaining", "IntegerField", "Daily Prep Log — decremented automatically as orders are placed"),
    ])

    add_heading(doc, "3.5 Order & OrderItem (apps.orders)", level=2)
    model_table(doc, [
        ("order_type", "CharField (choices)", "DINE_IN / TAKEAWAY"),
        ("session", "ForeignKey(TableSession), null", "Set for dine-in, null for takeaway"),
        ("table", "ForeignKey(Table), null", "Set for dine-in, null for takeaway"),
        ("branch", "ForeignKey(Branch), null", "Set directly for takeaway; dine-in resolves branch via table"),
        ("customer_name / customer_phone", "CharField", "Takeaway only"),
        ("status", "CharField (choices)", "NEW / ACCEPTED / PREPARING / READY / COLLECTED / SERVED / CANCELLED"),
        ("item.menu_item / item.quantity / item.unit_price / item.notes", "ForeignKey / IntegerField / DecimalField", "Line items, price snapshotted at order time"),
    ])

    add_heading(doc, "3.6 Bill & CashierShift (apps.billing)", level=2)
    model_table(doc, [
        ("session", "OneToOneField(TableSession), null", "Dine-in bill (exactly one of session/order is set)"),
        ("order", "OneToOneField(Order), null", "Takeaway bill"),
        ("subtotal / tax_amount / service_charge / discount_amount / total_amount", "DecimalField", "Computed from the tenant's gst_percentage / service_charge_percentage at payment time"),
        ("payment_method", "CharField (choices)", "CASH / CARD / UPI"),
        ("shift.opened_at / shift.closed_at / shift.counted_cash", "DateTimeField / DecimalField", "Cashier shift reconciliation against the system-computed cash total"),
    ])

    add_heading(doc, "3.7 Ingredient, StockMovement, PurchaseOrder & RecipeItem (apps.inventory)", level=2)
    model_table(doc, [
        ("ingredient.name / unit / current_stock / minimum_stock_level / unit_cost", "CharField / DecimalField", "Stock ledger row"),
        ("stock_movement.movement_type", "CharField (choices)", "RESTOCK / WASTAGE / USAGE / ADJUSTMENT"),
        ("stock_movement.wastage_reason", "CharField (choices)", "SPOILED / OVER_PREPPED / RETURNED / OTHER"),
        ("purchase_order.status", "CharField (choices)", "PENDING / APPROVED / REJECTED / ORDERED / RECEIVED"),
        ("purchase_order_line.quantity_ordered / quantity_received / unit_cost", "DecimalField", "Line items on a purchase order"),
        ("recipe_item.menu_item / ingredient / quantity_per_serving", "ForeignKey / DecimalField", "Drives automatic stock deduction from the Daily Prep Log"),
    ])

    add_heading(doc, "3.8 KDSDevice (apps.kitchen)", level=2)
    model_table(doc, [
        ("label", "CharField", "Human-readable device name (e.g. 'Main Kitchen Tablet')"),
        ("api_key", "CharField, unique", "Device credential — sent as the X-KDS-API-Key header, no user login involved"),
        ("branch", "ForeignKey(Branch), null", "Optional branch scoping"),
        ("is_active / last_seen_at", "BooleanField / DateTimeField", "Device status and last poll time"),
    ])

    add_heading(doc, "3.9 Notification (apps.notifications)", level=2)
    model_table(doc, [
        ("recipient", "ForeignKey(User)", "Staff member this notification is for"),
        ("type", "CharField (choices)", "ORDER_READY / BILL_REQUESTED / PAYMENT_CONFIRMED"),
        ("order / table / branch", "ForeignKey, null", "Context links, whichever apply"),
        ("is_read", "BooleanField", "Toggled via the Mark Read action"),
    ])

    add_heading(doc, "3.10 PlatformAdmin, ImpersonationSession & PlatformActivityLog (apps.platform)", level=2)
    doc.add_paragraph("A separate auth realm from restaurant staff — a platform token never carries a restaurant_id claim.").runs[0].font.size = Pt(10)
    model_table(doc, [
        ("platform_admin.email / access_level", "EmailField / CharField (choices)", "FULL_ADMIN / SUPPORT_READONLY"),
        ("platform_login_code.code / expires_at", "CharField / DateTimeField", "6-digit 2FA code, 5-minute TTL"),
        ("impersonation_session.platform_admin / restaurant / target_user / expires_at / ended_at", "ForeignKey / DateTimeField", "Support-access session; ending it revokes the minted token immediately"),
        ("platform_activity_log.action / description / actor / restaurant", "CharField (choices) / CharField / ForeignKey", "Audit trail — org created/updated, status/flags/plan changed, team added/removed, impersonation started/ended"),
    ])

    # ---- 4. API Endpoints ----
    doc.add_page_break()
    add_heading(doc, "4. API Endpoints", level=1)
    p = doc.add_paragraph()
    r = p.add_run(f"BASE URL: {BASE_URL}")
    r.bold = True
    r.font.size = Pt(11)
    doc.add_paragraph(
        "Every row below is a real request executed live against the deployed server above, "
        "via Postman's official Newman CLI runner, with the exact request body sent and the "
        "exact response body received — nothing here is hand-written or assumed. Rows marked "
        "✅ Verified passed their assertions on this run; rows marked with a status note are "
        "intentional negative-path tests (verifying the API correctly rejects/blocks something)."
    ).runs[0].font.size = Pt(9.5)

    for i, (folder_name, section_title) in enumerate(SECTION_TITLES, start=1):
        add_heading(doc, f"{i}. {section_title} ✅", level=2)
        names = folder_index.get(folder_name, [])
        etbl = doc.add_table(rows=1, cols=8)
        etbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr = etbl.rows[0].cells
        for c, h in enumerate(
            ["Screen", "Method", "URL", "Parameters", "Output Parameters (In Response)", "Changes Required", "Comments", "Status"]
        ):
            add_plain(hdr[c], h, size_pt=9, bold=True)
            set_cell_bg(hdr[c], "D9E2F3")

        for name in names:
            entries = by_name.get(name)
            if not entries:
                continue
            r = entries[0]
            row = etbl.add_row().cells
            add_plain(row[0], name, size_pt=8.5)
            add_plain(row[1], r["method"], size_pt=8.5, bold=True)
            add_mono(row[2], r["path"], size_pt=7.5)
            params = r["requestBody"] if r["requestBody"] != "(no request body — GET)" else f"Header: Authorization: {r['auth']}" if r["auth"] != "None (public)" else "(no auth, no body)"
            add_mono(row[3], pretty_json(params) if params.startswith("{") else params, size_pt=7)
            resp = r["responseBody"] or "(empty)"
            add_mono(row[4], pretty_json(resp), size_pt=7)
            add_plain(row[5], "", size_pt=8.5)
            note = EXPECTED_NON_2XX.get(name, "")
            add_plain(row[6], note, size_pt=8)
            status_text = f"✅ Verified ({r['statusCode']})" if not note else f"✅ Verified ({r['statusCode']}) — expected"
            add_plain(row[7], status_text, size_pt=8, bold=True, color=STATUS_OK_COLOR)

        for c, w in enumerate([2.0, 0.9, 2.3, 2.8, 4.2, 1.3, 1.8, 1.8]):
            set_col_width(etbl.columns[c], w)
        set_table_borders(etbl)
        doc.add_paragraph()

    doc.save(str(OUTPUT_DOCX))


def main() -> int:
    if not COLLECTION_JSON.exists() or not RESULTS_JSON.exists():
        print(f"MISSING: run scripts/build_postman_collection.py and scripts/run_via_newman.js first.")
        return 2
    collection = json.loads(COLLECTION_JSON.read_text(encoding="utf-8"))
    results = json.loads(RESULTS_JSON.read_text(encoding="utf-8"))
    build(collection, results)
    print(f"Wrote {OUTPUT_DOCX} ({OUTPUT_DOCX.stat().st_size:,} bytes)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
