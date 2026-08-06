"""
Generate a Word document (.docx) report for the 01 Auth folder Postman run.

Output: docs/DineOS_01_Auth_Test_Report.docx
"""

from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

# ---------------------------------------------------------------------------
# Test run data (unredacted, exactly as returned by the live server)
# ---------------------------------------------------------------------------

ENVIRONMENT = {
    "server_root": "https://dineos-1unt.onrender.com",
    "base_url": "https://dineos-1unt.onrender.com/v1",
    "tenant_slug": "demo-bistro",
    "collection": "DineOS v2 Backend API (Phase 1)",
    "folder": "01 Auth",
    "run_date_utc": "2026-08-04",
}

RESULTS = [
    {
        "step": 1,
        "name": "Login — Admin",
        "method": "POST",
        "path": "/v1/auth/login/",
        "full_url": "https://dineos-1unt.onrender.com/v1/auth/login/",
        "auth": "None (public)",
        "request_body": '{"email": "admin@demo-bistro.demo", "password": "Demo@1234"}',
        "response_body": (
            '{"refresh":"eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJ0b2tlbl90eXBlIjoi'
            "cmVmcmVzaCIsImV4cCI6MTc4NjQyNTc3MSwiaWF0IjoxNzg1ODIwOTcxLCJqdGkiOiIz"
            "NTQzNzU3NTcxN2M0NjE0ODcxOWI5MzcyYWU0MWJmMCIsInVzZXJfaWQiOiJhY2M1MDE0"
            "Zi0wOGU3LTRjM2EtYjJmYy0wY2EyMjk4Y2M2NzciLCJyb2xlIjoiQURNSU4iLCJuYW1l"
            "IjoiQWRtaW4iLCJyZXN0YXVyYW50X2lkIjoiMSJ9.384n9AQhwTVBQoYCm7l3K8d4aiq"
            'DjzVU-AXHzcqEKcM","access":"eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJ'
            "0b2tlbl90eXBlIjoiYWNjZXNzIiwiZXhwIjoxNzg1OTA3MzcxLCJpYXQiOjE3ODU4MjA"
            "5NzEsImp0aSI6IjM5ZDk2NjRkMDI2YTRlZDdhZDQ1NGQ0MTcxODY2MWEzIiwidXNlcl9"
            "pZCI6ImFjYzUwMTRmLTA4ZTctNGMzYS1iMmZjLTBjYTIyOThjYzY3NyIsInJvbGUiOiJ"
            "BRE1JTiIsIm5hbWUiOiJBZG1pbiIsInJlc3RhdXJhbnRfaWQiOiIxIn0.AUxmoelkd8S"
            'p34zi4QelMWPuymf-mZmoZIL5BPyXeCY","role":"ADMIN","name":"Admin",'
            '"restaurant_id":"1"}'
        ),
        "status_code": 200,
        "pass_fail": "PASS",
        "notes": "afterResponse script set collection vars admin_token and refresh_token.",
    },
    {
        "step": 2,
        "name": "Login — Manager",
        "method": "POST",
        "path": "/v1/auth/login/",
        "full_url": "https://dineos-1unt.onrender.com/v1/auth/login/",
        "auth": "None (public)",
        "request_body": '{"email": "manager@demo-bistro.demo", "password": "Demo@1234"}',
        "response_body": (
            '{"refresh":"eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJ0b2tlbl90eXBlIjoi'
            "cmVmcmVzaCIsImV4cCI6MTc4NjQyNTc5MywiaWF0IjoxNzg1ODIwOTkzLCJqdGkiOiIz"
            "MDUxOTYxMGI5YWM0OTBjYTVhYjQ4MjdkYjc2N2U2NSIsInVzZXJfaWQiOiJkMzE4ZDM4"
            "Zi1mMmI3LTQ1MGMtOGIyZS0wMjEwZTkyOGFiNzMiLCJyb2xlIjoiTUFOQUdFUiIsIm5h"
            'bWUiOiJNYW5hZ2VyIiwicmVzdGF1cmFudF9pZCI6IjEifQ.yOFvA6chP0oBOyaBgD8VYt'
            'yWeBO0LubQ5ahkU73pNjc","access":"eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVC'
            "J9.eyJ0b2tlbl90eXBlIjoiYWNjZXNzIiwiZXhwIjoxNzg1OTA3MzkzLCJpYXQiOjE3O"
            "DU4MjA5OTMsImp0aSI6IjA1Yzk5YjkwMDM4ZDRmMzQ5NjhmNjgzOGFhZTQzNzAxIiwid"
            "XNlcl9pZCI6ImQzMThkMzhmLWYyYjctNDUwYy04YjJlLTAyMTBlOTI4YWI3MyIsInJvb"
            "GUiOiJNQU5BR0VSIiwibmFtZSI6Ik1hbmFnZXIiLCJyZXN0YXVyYW50X2lkIjoiMSJ9."
            'xGJSf_f07e662TJ1zO2B56K6UtnEPlVAipSbj5oOD4w","role":"MANAGER","name'
            '":"Manager","restaurant_id":"1"}'
        ),
        "status_code": 200,
        "pass_fail": "PASS",
        "notes": "afterResponse script set collection var manager_token.",
    },
    {
        "step": 3,
        "name": "Login — Server",
        "method": "POST",
        "path": "/v1/auth/login/",
        "full_url": "https://dineos-1unt.onrender.com/v1/auth/login/",
        "auth": "None (public)",
        "request_body": '{"email": "server@demo-bistro.demo", "password": "Demo@1234"}',
        "response_body": (
            '{"refresh":"eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJ0b2tlbl90eXBlIjoi'
            "cmVmcmVzaCIsImV4cCI6MTc4NjQyNTgxNCwiaWF0IjoxNzg1ODIxMDE0LCJqdGkiOiI5"
            "NDFlZGMxYjI1N2M0ZTBlYWU1NDY3NjAyYjA2MGQzMyIsInVzZXJfaWQiOiJlYmJmMDk1"
            "Yy05NzBlLTRkOWYtYTZhNS1mZjBkNGE4ZTg3MzgiLCJyb2xlIjoiU0VSVkVSIiwibmFt"
            'ZSI6IlNlcnZlciIsInJlc3RhdXJhbnRfaWQiOiIxIn0.7n_jKDUIcHaeTLFKjYgX6KVKX'
            'Frxex0epsCIMQDoS5I","access":"eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.'
            "eyJ0b2tlbl90eXBlIjoiYWNjZXNzIiwiZXhwIjoxNzg1OTA3NDE0LCJpYXQiOjE3ODU4"
            "MjEwMTQsImp0aSI6ImY4MmQ1NzEyZDg1YjRkODNhMGU3OTMxYzQwMTVkNzdmIiwidXNl"
            "cl9pZCI6ImViYmYwOTVjLTk3MGUtNGQ5Zi1hNmE1LWZmMGQ0YThlODczOCIsInJvbGUi"
            "OiJTRVJWRVIiLCJuYW1lIjoiU2VydmVyIiwicmVzdGF1cmFudF9pZCI6IjEifQ.LM0DX"
            'TY4TxXKMpdkLj0RkzqVn4khICsDeQMaG5T3SmE","role":"SERVER","name":"Se'
            'rver","restaurant_id":"1"}'
        ),
        "status_code": 200,
        "pass_fail": "PASS",
        "notes": "afterResponse script set collection var server_token.",
    },
    {
        "step": 4,
        "name": "Login — Cashier",
        "method": "POST",
        "path": "/v1/auth/login/",
        "full_url": "https://dineos-1unt.onrender.com/v1/auth/login/",
        "auth": "None (public)",
        "request_body": '{"email": "cashier@demo-bistro.demo", "password": "Demo@1234"}',
        "response_body": (
            '{"refresh":"eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJ0b2tlbl90eXBlIjoi'
            "cmVmcmVzaCIsImV4cCI6MTc4NjQyNTgzMiwiaWF0IjoxNzg1ODIxMDMyLCJqdGkiOiI0"
            "N2EwZDMzOTcwYWM0ZjlkODE5MDg3YzYyMTk2ZmJlZSIsInVzZXJfaWQiOiJkNDExZTVj"
            "OC1jMmMyLTQ1YzUtOTI2OC1kOTMwNzU4NzgxOGIiLCJyb2xlIjoiQ0FTSElFUiIsIm5h"
            'bWUiOiJDYXNoaWVyIiwicmVzdGF1cmFudF9pZCI6IjEifQ.LFZGWLqsK1p12eik3Legg'
            '9x3bj8jMYTkzhfDrk4IEcw","access":"eyJhbGciOiJIUzI1NiIsInR5cCI6IkpX'
            "VCJ9.eyJ0b2tlbl90eXBlIjoiYWNjZXNzIiwiZXhwIjoxNzg1OTA3NDMyLCJpYXQiOjE"
            "3ODU4MjEwMzIsImp0aSI6ImFmMDJjNjk1MDczYTQ2NTVhM2Y2MjhkMWMzODUxMGNhIiw"
            "idXNlcl9pZCI6ImQ0MTFlNWM4LWMyYzItNDVjNS05MjY4LWQ5MzA3NTg3ODE4YiIsInJ"
            "vbGUiOiJDQVNISUVSIiwibmFtZSI6IkNhc2hpZXIiLCJyZXN0YXVyYW50X2lkIjoiMSJ"
            '9.7-O0CPvs2FWtdOrOgAViWkmbRG2UhVyQTiYk1W3FDq8","role":"CASHIER","na'
            'me":"Cashier","restaurant_id":"1"}'
        ),
        "status_code": 200,
        "pass_fail": "PASS",
        "notes": "afterResponse script set collection var cashier_token.",
    },
    {
        "step": 5,
        "name": "Me",
        "method": "GET",
        "path": "/v1/auth/me/",
        "full_url": "https://dineos-1unt.onrender.com/v1/auth/me/",
        "auth": "Bearer {{admin_token}}",
        "request_body": "(no request body — GET)",
        "response_body": (
            '{"id":"acc5014f-08e7-4c3a-b2fc-0ca2298cc677","email":"admin@demo-bistro.demo",'
            '"name":"Admin","phone":"","role":"ADMIN","is_active":true,'
            '"restaurant":{"id":1,"name":"DineOS Demo Bistro","slug":"demo-bistro",'
            '"notifications_enabled":true,"kitchen_enabled":true,"billing_enabled":true,'
            '"realtime_enabled":true},"created_at":"2026-08-03T12:04:47.278027Z"}'
        ),
        "status_code": 200,
        "pass_fail": "PASS",
        "notes": "Inline test asserts all four tenant feature flags present — passed.",
    },
    {
        "step": 6,
        "name": "Refresh Token",
        "method": "POST",
        "path": "/v1/auth/refresh-token/",
        "full_url": "https://dineos-1unt.onrender.com/v1/auth/refresh-token/",
        "auth": "None (refresh JWT is in body)",
        "request_body": (
            '{"refresh": "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJ0b2tlbl90eXBlIjoi'
            "cmVmcmVzaCIsImV4cCI6MTc4NjQyNTc3MSwiaWF0IjoxNzg1ODIwOTcxLCJqdGkiOiIz"
            "NTQzNzU3NTcxN2M0NjE0ODcxOWI5MzcyYWU0MWJmMCIsInVzZXJfaWQiOiJhY2M1MDE0"
            "Zi0wOGU3LTRjM2EtYjJmYy0wY2EyMjk4Y2M2NzciLCJyb2xlIjoiQURNSU4iLCJuYW1l"
            'IjoiQWRtaW4iLCJyZXN0YXVyYW50X2lkIjoiMSJ9.384n9AQhwTVBQoYCm7l3K8d4aiq'
            'DjzVU-AXHzcqEKcM"}'
        ),
        "response_body": (
            '{"access":"eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJ0b2tlbl90eXBlIjoi'
            "YWNjZXNzIiwiZXhwIjoxNzg1OTA3NDYxLCJpYXQiOjE3ODU4MjA5NzEsImp0aSI6IjUx"
            "ZGNlMzYwMzgzZTQ5MDdiNjM4NzZjODI0YjljZDNmIiwidXNlcl9pZCI6ImFjYzUwMTRm"
            "LTA4ZTctNGMzYS1iMmZjLTBjYTIyOThjYzY3NyIsInJvbGUiOiJBRE1JTiIsIm5hbWUi"
            'OiJBZG1pbiIsInJlc3RhdXJhbnRfaWQiOiIxIn0.6mfIqoHVY80ibtzVt89YrKilLd94'
            'NtT-bNRomhM5Wck","refresh":"eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.ey'
            "J0b2tlbl90eXBlIjoicmVmcmVzaCIsImV4cCI6MTc4NjQyNTg2MSwiaWF0IjoxNzg1OD"
            "IxMDYxLCJqdGkiOiJkYmMyYjRkNjJiODA0ZTBjODA1YjQyMzQ3MGJmYjZlNiIsInVzZX"
            "JfaWQiOiJhY2M1MDE0Zi0wOGU3LTRjM2EtYjJmYy0wY2EyMjk4Y2M2NzciLCJyb2xlIj"
            "oiQURNSU4iLCJuYW1lIjoiQWRtaW4iLCJyZXN0YXVyYW50X2lkIjoiMSJ9.pSjaK_zS-"
            'FdvTB6s8uWnygepcV3vNJdwQ0Q4FL5lNR4"}'
        ),
        "status_code": 200,
        "pass_fail": "PASS",
        "notes": (
            "SimpleJWT rotated the refresh token (new refresh JWT returned). "
            "The Refresh Token request has no afterResponse script, so the collection "
            "variable refresh_token was NOT updated by this call."
        ),
    },
    {
        "step": 7,
        "name": "Change Password",
        "method": "PATCH",
        "path": "/v1/auth/change-password/",
        "full_url": "https://dineos-1unt.onrender.com/v1/auth/change-password/",
        "auth": "Bearer {{admin_token}}",
        "request_body": '{"current_password": "Demo@1234", "new_password": "Demo@1234"}',
        "response_body": '{"detail":"Password changed."}',
        "status_code": 200,
        "pass_fail": "PASS",
        "notes": "new_password intentionally same as current so demo account remains usable.",
    },
    {
        "step": 8,
        "name": "Forgot Password",
        "method": "POST",
        "path": "/v1/auth/forgot-password/",
        "full_url": "https://dineos-1unt.onrender.com/v1/auth/forgot-password/",
        "auth": "None (public)",
        "request_body": '{"email": "admin@demo-bistro.demo"}',
        "response_body": (
            '{"detail":"If that email is registered, a reset link has been sent."}'
        ),
        "status_code": 200,
        "pass_fail": "PASS",
        "notes": (
            "Token is only returned in the response body when Django DEBUG=True. "
            "On the deployed Render service DEBUG=False, so no reset_token was "
            "captured. The token is expected to be delivered via email in production."
        ),
    },
    {
        "step": 9,
        "name": "Logout",
        "method": "POST",
        "path": "/v1/auth/logout/",
        "full_url": "https://dineos-1unt.onrender.com/v1/auth/logout/",
        "auth": "Bearer {{admin_token}}",
        "request_body": (
            '{"refresh": "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJ0b2tlbl90eXBlIjoi'
            "cmVmcmVzaCIsImV4cCI6MTc4NjQyNTg2MSwiaWF0IjoxNzg1ODIxMDYxLCJqdGkiOiJk"
            "YmMyYjRkNjJiODA0ZTBjODA1YjQyMzQ3MGJmYjZlNiIsInVzZXJfaWQiOiJhY2M1MDE0"
            "Zi0wOGU3LTRjM2EtYjJmYy0wY2EyMjk4Y2M2NzciLCJyb2xlIjoiQURNSU4iLCJuYW1l"
            'IjoiQWRtaW4iLCJyZXN0YXVyYW50X2lkIjoiMSJ9.pSjaK_zS-FdvTB6s8uWnygepcV3'
            'vNJdwQ0Q4FL5lNR4"}'
        ),
        "response_body": "(empty — 204 No Content)",
        "status_code": 204,
        "pass_fail": "PASS",
        "notes": "204 No Content — refresh token blacklisted server-side.",
    },
]

OBSERVATIONS = [
    (
        "Consistent login response shape",
        "All four role logins (Admin, Manager, Server, Cashier) return the same JSON "
        "structure — {refresh, access, role, name, restaurant_id} — with role and "
        "name reflecting the user, and restaurant_id: \"1\" for every user "
        "(DineOS Demo Bistro tenant).",
    ),
    (
        "/auth/me/ exposes tenant feature flags",
        "The Me endpoint returns the full user profile plus a nested `restaurant` "
        "object with kitchen_enabled, billing_enabled, notifications_enabled, and "
        "realtime_enabled booleans. All four are true on the demo-bistro tenant. "
        "The inline collection test asserting these four fields passed.",
    ),
    (
        "Refresh token rotates on refresh",
        "SimpleJWT is configured with ROTATE_REFRESH_TOKENS=True. The refresh-token "
        "endpoint returns BOTH a new access AND a new refresh JWT. Consumers must "
        "store the new refresh JWT and discard the old one. The Postman request "
        "currently has no afterResponse script, so the collection variable "
        "refresh_token is not updated automatically — consider adding one.",
    ),
    (
        "Forgot Password does not leak the token in production",
        "The endpoint returns only a generic acknowledgement message. The "
        "PasswordResetToken is only exposed in the response body when "
        "settings.DEBUG=True (see apps/authentication/views.py:60-61). In production "
        "the token must be delivered via email. This is correct security behavior.",
    ),
    (
        "Logout returns 204 No Content",
        "The logout endpoint blacklists the supplied refresh token server-side and "
        "returns 204 with no response body. Clients should treat any 2xx status as "
        "success and not attempt to parse a body.",
    ),
]

# ---------------------------------------------------------------------------
# Docx helpers
# ---------------------------------------------------------------------------

STATUS_COLORS = {
    "PASS": RGBColor(0x0B, 0x71, 0x2A),   # dark green
    "FAIL": RGBColor(0xB3, 0x1E, 0x1E),   # dark red
}


def set_cell_bg(cell, hex_color: str) -> None:
    """Set a table cell's shading fill color."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_col_width(column, width_cm: float) -> None:
    for cell in column.cells:
        cell.width = Cm(width_cm)


def add_heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"


def add_mono(cell, text: str, size_pt: int = 8) -> None:
    """Fill a table cell with monospaced text, wrapping enabled."""
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(size_pt)
    # Also set east-asian font so line-break/wrap uses same face
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Consolas")
    rFonts.set(qn("w:hAnsi"), "Consolas")
    rFonts.set(qn("w:cs"), "Consolas")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def add_plain(cell, text: str, size_pt: int = 10, bold: bool = False,
              color: RGBColor | None = None) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def set_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "808080")
        tblBorders.append(b)
    tblPr.append(tblBorders)


# ---------------------------------------------------------------------------
# Build the document
# ---------------------------------------------------------------------------

def build_document(output_path: Path) -> None:
    doc = Document()

    # Margins: give tables more room
    for section in doc.sections:
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)

    # ---- Title page ----
    title = doc.add_heading("DineOS v2 Backend API — Phase 1", level=0)
    for run in title.runs:
        run.font.name = "Calibri"

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("01 Auth — Test Run Report")
    subtitle_run.font.size = Pt(16)
    subtitle_run.bold = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(6)
    meta.add_run(
        f"Collection: {ENVIRONMENT['collection']}\n"
        f"Folder: {ENVIRONMENT['folder']}\n"
        f"Target: {ENVIRONMENT['base_url']}\n"
        f"Tenant: {ENVIRONMENT['tenant_slug']}\n"
        f"Run date (UTC): {ENVIRONMENT['run_date_utc']}\n"
        f"Document generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    ).font.size = Pt(10)

    # ---- Summary paragraph ----
    add_heading(doc, "Overview", level=1)
    total = len(RESULTS)
    passed = sum(1 for r in RESULTS if r["pass_fail"] == "PASS")
    failed = total - passed
    p = doc.add_paragraph()
    p.add_run(
        f"All {total} requests in the 01 Auth folder were executed in order "
        f"against the live Render deployment at {ENVIRONMENT['base_url']}. "
        f"Result: {passed} passed, {failed} failed. Every 2xx response was saved "
        "as a Postman example alongside its request. This document captures the "
        "unredacted request and response bodies for each endpoint, intended as the "
        "source of truth for building the official API documentation."
    ).font.size = Pt(10.5)

    # ---- Summary table (compact) ----
    add_heading(doc, "Summary at a glance", level=1)
    summary_tbl = doc.add_table(rows=1, cols=5)
    summary_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = summary_tbl.rows[0].cells
    for i, h in enumerate(["#", "Endpoint", "Method", "Status", "Result"]):
        add_plain(hdr[i], h, size_pt=10, bold=True)
        set_cell_bg(hdr[i], "D9E2F3")

    for r in RESULTS:
        row = summary_tbl.add_row().cells
        add_plain(row[0], str(r["step"]), size_pt=10)
        add_plain(row[1], f"{r['name']}\n{r['path']}", size_pt=9)
        add_plain(row[2], r["method"], size_pt=10)
        add_plain(row[3], str(r["status_code"]), size_pt=10)
        add_plain(
            row[4], r["pass_fail"], size_pt=10, bold=True,
            color=STATUS_COLORS[r["pass_fail"]],
        )

    # Column widths
    widths_cm = [1.0, 6.5, 2.0, 2.0, 2.0]
    for col_idx, w in enumerate(widths_cm):
        set_col_width(summary_tbl.columns[col_idx], w)
    set_table_borders(summary_tbl)

    # ---- Full details table ----
    doc.add_page_break()
    add_heading(doc, "Full request/response details", level=1)
    doc.add_paragraph(
        "The tables below show the exact JSON that was sent to and received "
        "from the live server, with no truncation. JWT values are complete."
    ).runs[0].font.size = Pt(10)

    for r in RESULTS:
        add_heading(doc, f"{r['step']}. {r['name']}  —  {r['method']} {r['path']}",
                    level=2)

        detail_tbl = doc.add_table(rows=0, cols=2)
        detail_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        detail_tbl.autofit = False

        def add_kv(label: str, value: str, mono: bool = False):
            row = detail_tbl.add_row().cells
            add_plain(row[0], label, size_pt=10, bold=True)
            set_cell_bg(row[0], "F2F2F2")
            if mono:
                add_mono(row[1], value)
            else:
                add_plain(row[1], value, size_pt=10)

        add_kv("Method", r["method"])
        add_kv("Path", r["path"])
        add_kv("Full URL", r["full_url"], mono=True)
        add_kv("Auth", r["auth"])
        add_kv("Request Body", r["request_body"], mono=True)
        add_kv("Response Body", r["response_body"], mono=True)
        add_kv("Status Code", str(r["status_code"]))
        # Pass/fail with color
        row = detail_tbl.add_row().cells
        add_plain(row[0], "Result", size_pt=10, bold=True)
        set_cell_bg(row[0], "F2F2F2")
        add_plain(row[1], r["pass_fail"], size_pt=10, bold=True,
                  color=STATUS_COLORS[r["pass_fail"]])
        add_kv("Notes", r["notes"] or "—")

        # Widths for detail table
        set_col_width(detail_tbl.columns[0], 3.5)
        set_col_width(detail_tbl.columns[1], 14.0)
        set_table_borders(detail_tbl)

        doc.add_paragraph()  # spacer

    # ---- Observations ----
    doc.add_page_break()
    add_heading(doc, "Notable observations", level=1)
    for title, body in OBSERVATIONS:
        add_heading(doc, title, level=2)
        p = doc.add_paragraph(body)
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(10.5)

    # ---- Files created ----
    add_heading(doc, "Files created / modified during this run", level=1)
    for line in [
        "Modified: postman/collections/DineOS v2 Backend API (Phase 1)/.resources/definition.yaml"
        " (server_root updated from http://127.0.0.1:8000 to https://dineos-1unt.onrender.com)",
        "Created: 9 example files under 01 Auth/.resources/<Request>.resources/examples/",
        "Modified: all 9 *.request.yaml files in 01 Auth/ to reference their new examples folder",
    ]:
        doc.add_paragraph(line, style="List Bullet").runs[0].font.size = Pt(10)

    doc.save(str(output_path))


if __name__ == "__main__":
    out = Path(__file__).resolve().parent.parent / "docs" / "DineOS_01_Auth_Test_Report.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    build_document(out)
    print(f"Wrote {out} ({out.stat().st_size:,} bytes)")
